#!/usr/bin/env python3
"""
Backend Test Suite for PDF Processing Endpoints
Tests all 5 core PDF tools: Merge, Split, Compress, PDF to Word, Word to PDF
"""

import requests
import io
import os
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import zipfile

# Backend URL from environment
BACKEND_URL = "https://pdf-converter-168.preview.emergentagent.com/api"

class PDFTestSuite:
    def __init__(self):
        self.results = []
        self.sample_files = {}
        
    def log_result(self, test_name, success, message, details=None):
        """Log test result"""
        result = {
            "test": test_name,
            "success": success,
            "message": message,
            "details": details or {}
        }
        self.results.append(result)
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status}: {test_name} - {message}")
        if details:
            for key, value in details.items():
                print(f"  {key}: {value}")
        print()

    def create_sample_pdf(self, filename, content="Sample PDF content for testing"):
        """Create a sample PDF file"""
        buffer = io.BytesIO()
        pdf_canvas = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Add multiple pages for better testing
        for page_num in range(1, 4):  # 3 pages
            pdf_canvas.drawString(50, height - 50, f"Page {page_num}")
            pdf_canvas.drawString(50, height - 80, content)
            pdf_canvas.drawString(50, height - 110, f"This is page {page_num} of the test PDF")
            pdf_canvas.drawString(50, height - 140, "Generated for PDF processing tests")
            if page_num < 3:
                pdf_canvas.showPage()
        
        pdf_canvas.save()
        buffer.seek(0)
        
        self.sample_files[filename] = buffer.getvalue()
        return buffer.getvalue()

    def create_sample_docx(self, filename, content="Sample DOCX content for testing"):
        """Create a sample DOCX file"""
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph(content)
        doc.add_paragraph('This is a test DOCX file created for PDF conversion testing.')
        doc.add_paragraph('It contains multiple paragraphs to test the conversion process.')
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        self.sample_files[filename] = buffer.getvalue()
        return buffer.getvalue()

    def test_health_check(self):
        """Test the health check endpoint"""
        try:
            response = requests.get(f"{BACKEND_URL}/pdf/health", timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                self.log_result(
                    "Health Check",
                    True,
                    "PDF service is healthy",
                    {
                        "status_code": response.status_code,
                        "services": data.get("services", {})
                    }
                )
                return True
            else:
                self.log_result(
                    "Health Check",
                    False,
                    f"Health check failed with status {response.status_code}",
                    {"response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_result(
                "Health Check",
                False,
                f"Health check failed with error: {str(e)}"
            )
            return False

    def test_merge_pdf(self):
        """Test PDF merge endpoint"""
        try:
            # Create two sample PDFs
            pdf1 = self.create_sample_pdf("test1.pdf", "Content for first PDF")
            pdf2 = self.create_sample_pdf("test2.pdf", "Content for second PDF")
            
            files = [
                ('files', ('test1.pdf', io.BytesIO(pdf1), 'application/pdf')),
                ('files', ('test2.pdf', io.BytesIO(pdf2), 'application/pdf'))
            ]
            
            response = requests.post(f"{BACKEND_URL}/pdf/merge", files=files, timeout=60)
            
            if response.status_code == 200:
                # Check if response is a PDF
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                if 'application/pdf' in content_type and 'merged.pdf' in content_disposition:
                    self.log_result(
                        "Merge PDF",
                        True,
                        "Successfully merged PDFs",
                        {
                            "status_code": response.status_code,
                            "content_type": content_type,
                            "file_size": len(response.content),
                            "content_disposition": content_disposition
                        }
                    )
                    return True
                else:
                    self.log_result(
                        "Merge PDF",
                        False,
                        "Response is not a proper PDF file",
                        {
                            "content_type": content_type,
                            "content_disposition": content_disposition
                        }
                    )
                    return False
            else:
                self.log_result(
                    "Merge PDF",
                    False,
                    f"Merge failed with status {response.status_code}",
                    {"response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_result(
                "Merge PDF",
                False,
                f"Merge test failed with error: {str(e)}"
            )
            return False

    def test_split_pdf(self):
        """Test PDF split endpoint"""
        try:
            # Create a multi-page PDF
            pdf_content = self.create_sample_pdf("test_split.pdf", "Content for splitting test")
            
            files = {'file': ('test_split.pdf', io.BytesIO(pdf_content), 'application/pdf')}
            
            response = requests.post(f"{BACKEND_URL}/pdf/split", files=files, timeout=60)
            
            if response.status_code == 200:
                # Check if response is a ZIP file
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                if 'application/zip' in content_type and 'split_pages.zip' in content_disposition:
                    # Try to open the ZIP file to verify it contains PDFs
                    try:
                        zip_buffer = io.BytesIO(response.content)
                        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
                            file_list = zip_file.namelist()
                            pdf_files = [f for f in file_list if f.endswith('.pdf')]
                            
                            self.log_result(
                                "Split PDF",
                                True,
                                "Successfully split PDF into pages",
                                {
                                    "status_code": response.status_code,
                                    "content_type": content_type,
                                    "zip_size": len(response.content),
                                    "pdf_pages": len(pdf_files),
                                    "files_in_zip": file_list
                                }
                            )
                            return True
                    except zipfile.BadZipFile:
                        self.log_result(
                            "Split PDF",
                            False,
                            "Response is not a valid ZIP file",
                            {"content_type": content_type}
                        )
                        return False
                else:
                    self.log_result(
                        "Split PDF",
                        False,
                        "Response is not a proper ZIP file",
                        {
                            "content_type": content_type,
                            "content_disposition": content_disposition
                        }
                    )
                    return False
            else:
                self.log_result(
                    "Split PDF",
                    False,
                    f"Split failed with status {response.status_code}",
                    {"response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_result(
                "Split PDF",
                False,
                f"Split test failed with error: {str(e)}"
            )
            return False

    def test_compress_pdf(self):
        """Test PDF compression endpoint"""
        try:
            # Create a sample PDF
            pdf_content = self.create_sample_pdf("test_compress.pdf", "Content for compression test")
            original_size = len(pdf_content)
            
            files = {'file': ('test_compress.pdf', io.BytesIO(pdf_content), 'application/pdf')}
            data = {'quality': 'medium'}
            
            response = requests.post(f"{BACKEND_URL}/pdf/compress", files=files, data=data, timeout=60)
            
            if response.status_code == 200:
                # Check if response is a PDF
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                compression_ratio = response.headers.get('X-Compression-Ratio', 'N/A')
                
                if 'application/pdf' in content_type and 'compressed.pdf' in content_disposition:
                    compressed_size = len(response.content)
                    
                    self.log_result(
                        "Compress PDF",
                        True,
                        "Successfully compressed PDF",
                        {
                            "status_code": response.status_code,
                            "content_type": content_type,
                            "original_size": original_size,
                            "compressed_size": compressed_size,
                            "compression_ratio": compression_ratio,
                            "content_disposition": content_disposition
                        }
                    )
                    return True
                else:
                    self.log_result(
                        "Compress PDF",
                        False,
                        "Response is not a proper PDF file",
                        {
                            "content_type": content_type,
                            "content_disposition": content_disposition
                        }
                    )
                    return False
            else:
                self.log_result(
                    "Compress PDF",
                    False,
                    f"Compression failed with status {response.status_code}",
                    {"response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_result(
                "Compress PDF",
                False,
                f"Compression test failed with error: {str(e)}"
            )
            return False

    def test_pdf_to_word(self):
        """Test PDF to Word conversion endpoint"""
        try:
            # Create a sample PDF
            pdf_content = self.create_sample_pdf("test_pdf_to_word.pdf", "Content for PDF to Word conversion test")
            
            files = {'file': ('test_pdf_to_word.pdf', io.BytesIO(pdf_content), 'application/pdf')}
            
            response = requests.post(f"{BACKEND_URL}/pdf/pdf-to-word", files=files, timeout=60)
            
            if response.status_code == 200:
                # Check if response is a DOCX file
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                expected_content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                
                if expected_content_type in content_type and 'test_pdf_to_word.docx' in content_disposition:
                    self.log_result(
                        "PDF to Word",
                        True,
                        "Successfully converted PDF to Word",
                        {
                            "status_code": response.status_code,
                            "content_type": content_type,
                            "file_size": len(response.content),
                            "content_disposition": content_disposition
                        }
                    )
                    return True
                else:
                    self.log_result(
                        "PDF to Word",
                        False,
                        "Response is not a proper DOCX file",
                        {
                            "content_type": content_type,
                            "content_disposition": content_disposition,
                            "expected_content_type": expected_content_type
                        }
                    )
                    return False
            else:
                self.log_result(
                    "PDF to Word",
                    False,
                    f"PDF to Word conversion failed with status {response.status_code}",
                    {"response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_result(
                "PDF to Word",
                False,
                f"PDF to Word test failed with error: {str(e)}"
            )
            return False

    def test_word_to_pdf(self):
        """Test Word to PDF conversion endpoint"""
        try:
            # Create a sample DOCX
            docx_content = self.create_sample_docx("test_word_to_pdf.docx", "Content for Word to PDF conversion test")
            
            files = {'file': ('test_word_to_pdf.docx', io.BytesIO(docx_content), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            
            response = requests.post(f"{BACKEND_URL}/pdf/word-to-pdf", files=files, timeout=60)
            
            if response.status_code == 200:
                # Check if response is a PDF file
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                if 'application/pdf' in content_type and 'test_word_to_pdf.pdf' in content_disposition:
                    self.log_result(
                        "Word to PDF",
                        True,
                        "Successfully converted Word to PDF",
                        {
                            "status_code": response.status_code,
                            "content_type": content_type,
                            "file_size": len(response.content),
                            "content_disposition": content_disposition
                        }
                    )
                    return True
                else:
                    self.log_result(
                        "Word to PDF",
                        False,
                        "Response is not a proper PDF file",
                        {
                            "content_type": content_type,
                            "content_disposition": content_disposition
                        }
                    )
                    return False
            else:
                self.log_result(
                    "Word to PDF",
                    False,
                    f"Word to PDF conversion failed with status {response.status_code}",
                    {"response": response.text}
                )
                return False
                
        except Exception as e:
            self.log_result(
                "Word to PDF",
                False,
                f"Word to PDF test failed with error: {str(e)}"
            )
            return False

    def run_all_tests(self):
        """Run all PDF processing tests"""
        print("=" * 60)
        print("PDF PROCESSING ENDPOINTS TEST SUITE")
        print("=" * 60)
        print(f"Backend URL: {BACKEND_URL}")
        print()
        
        # Test order: Health check first, then all PDF operations
        tests = [
            self.test_health_check,
            self.test_merge_pdf,
            self.test_split_pdf,
            self.test_compress_pdf,
            self.test_pdf_to_word,
            self.test_word_to_pdf
        ]
        
        passed = 0
        total = len(tests)
        
        for test in tests:
            if test():
                passed += 1
        
        print("=" * 60)
        print("TEST SUMMARY")
        print("=" * 60)
        print(f"Total Tests: {total}")
        print(f"Passed: {passed}")
        print(f"Failed: {total - passed}")
        print(f"Success Rate: {(passed/total)*100:.1f}%")
        print()
        
        # Detailed results
        print("DETAILED RESULTS:")
        print("-" * 40)
        for result in self.results:
            status = "✅" if result["success"] else "❌"
            print(f"{status} {result['test']}: {result['message']}")
        
        return passed == total


if __name__ == "__main__":
    test_suite = PDFTestSuite()
    success = test_suite.run_all_tests()
    
    if success:
        print("\n🎉 All tests passed! PDF processing endpoints are working correctly.")
        exit(0)
    else:
        print("\n⚠️  Some tests failed. Check the detailed results above.")
        exit(1)